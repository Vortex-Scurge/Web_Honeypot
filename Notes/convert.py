from docx import Document
import re
import os

def clean_markdown(text):
    # Remove bold, italics, code backticks, and markdown links
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Remove horizontal rules
    if text.strip() in ['---', '***', '___']:
        return None
    return text.strip()

def md_to_docx(md_path, docx_path, title):
    if not os.path.exists(md_path):
        print(f"Skipping {md_path}, file not found.")
        return

    doc = Document()
    doc.add_heading(title, 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            doc.add_paragraph(line.rstrip(), style='No Spacing')
            continue

        if not stripped:
            continue
            
        # Handle headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        else:
            clean = clean_markdown(line)
            if clean:
                doc.add_paragraph(clean)
                
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == '__main__':
    md_to_docx('VIVA_PREP.md', 'VIVA_PREP.docx', 'Web Honeypot - VIVA PREPARATION')
    md_to_docx('CODE_EXPLANATION.md', 'CODE_EXPLANATION.docx', 'Web Honeypot - DETAILED CODE EXPLANATION')
    md_to_docx('web_honeypot_report.md', 'web_honeypot_report.docx', 'Web Honeypot - PROJECT REPORT')
